# create_lab_brochure.py

import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

# --- Helper function to add a bold heading ---
def add_section_heading(document, text):
    """Adds a bold, centered heading for a new section."""
    p = document.add_paragraph()
    p.add_run(text).bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)

# --- All content for the "Diagnostic & Laboratory Solutions" brochure ---

# Hero Section Content
HERO_TITLE = "Diagnostic & Laboratory Solutions"
HERO_SUBTITLE = "Automated Ecosystems for Accurate and Efficient Results"
HERO_DESCRIPTION = (
    "Complete laboratory ecosystems featuring automated chemistry analyzers, hematology systems, "
    "and comprehensive IVD solutions tailored for Ethiopian healthcare facilities."
)
IMAGE_PLACEHOLDER_NAME = "Scientist in modern lab with analysis machines" 

# On the website, these are called "Key Products", we'll list them as Key Features for brochure consistency
KEY_FEATURES = [
    "Mindray BS-240 Chemistry Analyzers",
    "Mindray BC-5150 Hematology Systems",
    "High-Precision Digital Microscopes",
    "Comprehensive Lab Furniture Solutions",
    "Integrated Quality Control Systems & Reagents",
]

# Benefits Content
BENEFITS = [
    "Increase patient throughput with automated workflows",
    "Reduce diagnostic errors by 85% with reliable results",
    "Accelerate staff competency with comprehensive training",
    "Minimize equipment downtime with 24/7 technical support",
    "Ensure consistent and traceable results",
]

# The website doesn't show tech specs for this section, so we'll create some representative ones as a template.
TECH_SPECS_DATA = [
    ("Chemistry Analyzer:", "Up to 240 tests/hour throughput", "Hematology System:", "5-part differentiation, 60 samples/hour"),
    ("Sample Types:", "Serum, plasma, urine, whole blood", "Data Management:", "LIS/HIS connectivity with HL7 protocol"),
    ("Reagent System:", "Refrigerated reagent positions, barcode reading", "Optical System:", "Advanced LED illumination microscopes"),
    ("Installation:", "Full setup, calibration, and method validation", "Support:", "On-site user training and remote diagnostics")
]

# --- Main function to build the Word document ---

def create_brochure():
    """Generates the brochure Word document from the predefined content."""
    document = Document()
    
    # Set document margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 1. Main Title and Subtitle Section
    document.add_heading(HERO_TITLE, level=1).alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_subtitle = document.add_paragraph()
    p_subtitle.add_run(HERO_SUBTITLE).bold = True
    p_subtitle.runs[0].font.color.rgb = RGBColor(0x1d, 0x8a, 0x5a) # Brand Green Color

    # 2. Main content block (Description and Image Placeholder)
    hero_table = document.add_table(rows=1, cols=2)
    hero_table.autofit = False
    hero_table.columns[0].width = Inches(4.0)
    hero_table.columns[1].width = Inches(2.5)
    
    # Left cell with description and call-to-action buttons
    left_cell = hero_table.cell(0, 0)
    left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    left_cell.text = HERO_DESCRIPTION
    
    p_buttons = left_cell.add_paragraph()
    p_buttons.paragraph_format.space_before = Pt(12)
    p_buttons.add_run("📄 Download Brochure").bold = True
    p_buttons.add_run("     ")
    p_buttons.add_run("🗓️ Schedule Demo").bold = True

    # Right cell with the image placeholder
    right_cell = hero_table.cell(0, 1)
    right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p_placeholder = right_cell.paragraphs[0]
    p_placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_placeholder = p_placeholder.add_run(f"[Image Placeholder:\n{IMAGE_PLACEHOLDER_NAME}]")
    font = run_placeholder.font
    font.italic = True
    font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # 3. Features and Benefits Section
    add_section_heading(document, "Key Features & Benefits")
    fb_table = document.add_table(rows=1, cols=2)
    fb_table.columns[0].width = Inches(3.25)
    fb_table.columns[1].width = Inches(3.25)

    # Populate Features
    features_cell = fb_table.cell(0, 0)
    features_cell.add_paragraph().add_run("Key Features").bold = True
    for item in KEY_FEATURES:
        features_cell.add_paragraph(item, style='List Bullet')

    # Populate Benefits
    benefits_cell = fb_table.cell(0, 1)
    benefits_cell.add_paragraph().add_run("Benefits").bold = True
    for item in BENEFITS:
        benefits_cell.add_paragraph(item, style='List Bullet')

    # 4. Technical Specifications Section
    add_section_heading(document, "Technical Specifications")
    
    spec_table = document.add_table(rows=len(TECH_SPECS_DATA), cols=4)
    spec_table.style = 'Table Grid'
    spec_table.autofit = False

    widths = [Inches(1.2), Inches(2.1), Inches(1.2), Inches(2.1)]
    for i, width in enumerate(widths):
        for cell in spec_table.columns[i].cells:
            cell.width = width
            
    # Populate the table with spec data
    for i, row_data in enumerate(TECH_SPECS_DATA):
        row_cells = spec_table.rows[i].cells
        for j, cell_text in enumerate(row_data):
            p = row_cells[j].paragraphs[0]
            runner = p.add_run(cell_text)
            if j % 2 == 0:
                runner.bold = True
    
    # --- Final step: Save the document ---
    file_name = "AteWork_Pharma_Lab_Solutions_Brochure.docx"
    document.save(file_name)
    print(f"✅ Brochure successfully generated: {file_name}")

# --- Run the main function ---
if __name__ == '__main__':
    try:
        create_brochure()
    except Exception as e:
        print(f"An unexpected error occurred: {e}")